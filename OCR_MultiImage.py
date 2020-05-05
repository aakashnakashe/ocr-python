try:
    from PIL import Image
except ImportError:
    import Image
import pytesseract as tess
tess.pytesseract.tesseract_cmd = r'.\tessdata\tesseract.exe'  #calling tessaract to env_path
from docx import Document
import os
import time
# import subprocess


userfilelocation = str(input('Your Folder Location: '))
useroutputlocation = str(input('Where do you want to store your Output: '))
os.chdir(userfilelocation)
print("\nImage File is scanning...")
time.sleep(3)
print("\nLooking for the data to extract...")
time.sleep(2)
print("\nDo you want to save this file in docx format")
answer = input('\nPress y for Yes and n for No: ')

try:
    for imageName in os.listdir(userfilelocation):
        if imageName.endswith('.png'):
            inputPath = os.path.join(userfilelocation, imageName)
            image = Image.open(inputPath)  #Enter image name  
            
            result = tess.image_to_string(image)
            
            fullTempPath = os.path.join(useroutputlocation, 'text_'+imageName+".txt") 
            # print(result)
        	# saving the text for every image in a separate .txt file 
            file1 = open(fullTempPath, "w") 
            file1.write(result) 
            file1.close()     
            
            # for saving the file in docx
            if answer == "y":            
                document = Document()
                myfile = open('text_'+imageName+'.txt').read()
                p = document.add_paragraph(myfile)
                document.save("word_"+imageName+".docx")
                print("\nYour word file has been saved successfully")
            else:        
                print("\nYour data is extracted from PNG Image")
        elif imageName.endswith('.jpg'):
            inputPath = os.path.join(userfilelocation, imageName)
            image = Image.open(inputPath)  #Enter image name  
            
            result = tess.image_to_string(image)
            
            fullTempPath = os.path.join(useroutputlocation, 'text_'+imageName+".txt") 
            # print(result)
        	# saving the text for every image in a separate .txt file 
            file1 = open(fullTempPath, "w") 
            file1.write(result) 
            file1.close()     
            
        #     for saving the file in docx
            if answer == "y":            
                document = Document()
                myfile = open('text_'+imageName+".txt").read()
                p = document.add_paragraph(myfile)
                document.save("word_"+imageName+".docx")
                print("\nYour word file has been saved successfully")
            else:        
                print("\nYour data is extracted from JPEG Image")
except Exception as error:
    print('\nCaught this error: ' + repr(error))

input('\nPress Enter to Continue...')
# subprocess.Popen(r'explorer /select, "useroutputlocation"')

os.execl(sys.executable, sys.executable, *sys.argv)
